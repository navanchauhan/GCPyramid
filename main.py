import pandas as pd
import tkinter as tk
import textwrap
from tkinter import ttk
from tkinter.filedialog import askopenfilename, asksaveasfilename
from datetime import datetime

from ttkwidgets import ScrolledListbox

import sv_ttk
import darkdetect

import os
import json
import sys

import subprocess
import platform

from os import path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from platformdirs import user_data_dir
from pathlib import Path

canadian_companies = []

# Config
appname = "PyramidScheme"
appauthor = "Navan Chauhan"
Path(user_data_dir(appname, appauthor)).mkdir( parents=True, exist_ok=True )

def get_companies_for_word_doc():
    None

def read_config(config_file='config.json'):
    default_config = {
        'img_width': 3300,
        'pyramid_width': 3000,
        'img_height': 1550,
        'overall_height': 2550,
        'block_color': [0, 150, 255],
        'dividend_color': [0, 150, 255],
        'font_color': [255, 255, 255],
        'normal_block': [0, 150, 255],
        'padding': 15,
        'radius': 10,
        'logo_padding': 100,
        'pyramid_padding_bottom': 350,
        'max_companies_in_row': 10,
    }
    
    if not os.path.exists(config_file):
        with open(config_file, 'w') as f:
            json.dump(default_config, f)
        return default_config
    else:
        with open(config_file, 'r') as f:
            config = json.load(f)
        return config

def pyramid_list(lst, sort_canadian=True):
    old_lst = lst

    top = []
    bottom = []

    for idx, company in enumerate(lst):
        if idx==0:
            top.append(company)
        else:
            if company in canadian_companies:
                bottom.append(company)
            else:
                top.append(company)

    lst = top + bottom

    i = 0
    rows = 0
    pyramid = []

    # Figure out the number of rows
    while i < len(lst):
        rows += 1
        i += rows

    # Generate the pyramid
    i = 0
    for r in range(1, rows+1):
        row = []
        for j in range(r):
            if i < len(lst):
                row.append(lst[i])
                i += 1
        pyramid.append(row)

    if len(pyramid) >= 2:
        if len(pyramid[-1]) < len(pyramid[-2]):
            pyramid[-2].extend(pyramid[-1])
            pyramid.pop()

    print("Checking for canadian companies")
    print(canadian_companies)
    if len(pyramid) > 2:
        print("Pyramid has more than two rows, checking to ensure Canadian companies are at the bottom")
        print(len(pyramid[-1]) - len(pyramid[-2]))
        for row in pyramid:
            print(len(row))
            print(row)
        if ((len(pyramid[-1]) - len(pyramid[-2])) > 1):
            print("Last two rows can be balanced if needed")
            for _ in range((len(pyramid[-1]) - len(pyramid[-2]))-1):
                company_to_add = None
                company_idx = 0
                for idx, company in enumerate(pyramid[-1]):
                    print(f'Checking {company}')
                    if company not in canadian_companies:
                        company_to_add = company
                        company_idx = idx
                        print(f'Shifting {company_to_add}')
                if company_to_add is not None:
                    pyramid[-1].pop(company_idx)
                    pyramid[-2].append(company_to_add)               

    for _ in range(4):
        if len(pyramid) > 2:
            if len(pyramid[-1]) < len(pyramid[-2]):
                pyramid[-3].append(pyramid[-2][-1])
                pyramid[-2].pop(-1)

    for _ in range(3):
        if len(pyramid) > 4:
            if len(pyramid[-2]) < len(pyramid[-3]):
                pyramid[-4].append(pyramid[-3][0])
                pyramid[-3].pop(0)

    return pyramid

class CompanySelector:
    def __init__(self, master):
        self.master = master
        self.master.title("Company Selector")
        self.master.geometry("600x600")
        self.select_file_button = ttk.Button(self.master, text="Select Excel File", command=self.load_file)
        self.select_file_button.pack(pady=250)
        print(path.join(user_data_dir(appname, appauthor), "default.txt"))
        # try:
        #     with open(path.join(user_data_dir(appname, appauthor), "default.txt")) as f:
        #         try:
        #             fname = f.read()
        #             self.df = pd.read_excel(fname)
        #             self.df = self.df[self.df['Company Name'].notna()]
        #             self.df = self.df[self.df['Symbol'].notna()]
        #             self.show_companies()
        #         except Exception as e:
        #             print(f'Oh No {e}')
        #             os.remove(path.join(user_data_dir(appname, appauthor), "default.txt"))
        #             self.load_file()
        # except FileNotFoundError:
        #     None
        self.show_companies()

    def load_file(self):
        file_path = askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
        if not file_path:
            return

        with open(path.join(user_data_dir(appname, appauthor), "default.txt"), "w") as f:
            f.write(file_path)
        try:
            self.df = pd.read_excel(file_path, keep_default_na=False)
            self.df = self.df[self.df['Company Name'].notna()]
            self.df = self.df[self.df['Symbol'].notna()]
        except Exception as e:
            print(f":( {e}")
            tk.messagebox.showerror(title="Whoops :(", message=f"{e}")
            return False
        return True
        #self.show_companies()


    def show_companies(self):
        try:
            with open(path.join(user_data_dir(appname, appauthor), "default.txt")) as f:
                try:
                    fname = f.read()
                    self.df = pd.read_excel(fname, keep_default_na=False)
                    self.df = self.df[self.df['Company Name'].notna()]
                    self.df = self.df[self.df['Symbol'].notna()]
                except Exception as e:
                    print(f'Oh No {e}')
                    tk.messagebox.showerror(title="Whoops! An Error Occurred", message=f"{e}\n\nYou will be asked to load the spreadsheet")
                    os.remove(path.join(user_data_dir(appname, appauthor), "default.txt"))
                    loaded_file = self.load_file()
                    while not loaded_file:
                        loaded_file = self.load_file()
        except FileNotFoundError:
            print("No default file set")
            tk.messagebox.showinfo(title="Hello world!", message="Since this is the first time you are using this software, please load the desired spreadsheet...")
            loaded_file = self.load_file()
            while not loaded_file:
                loaded_file = self.load_file()

        self.select_file_button.pack_forget()
        self.master.geometry("")
        self.master.resizable(width=False, height=False)

        width=600
        height=500
        screenwidth = self.master.winfo_screenwidth()
        screenheight = self.master.winfo_screenheight()
        alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
        self.master.geometry(alignstr)
        self.master.title("UCS Pyramid Builder")


        self.company_listbox = ScrolledListbox(self.master, selectmode=tk.MULTIPLE, exportselection=False, height=20)
        self.company_listbox.place(x=10,y=20,width=235,height=457)
        #self.company_listbox.pack(expand=True)

        for ticker, company in zip(self.df["Company Name"],self.df["Symbol"]):
            self.company_listbox.listbox.insert(tk.END, f' {company} - {ticker}')


        # bind the listbox to an onselect event
        self.company_listbox.listbox.bind('<<ListboxSelect>>', self.onselect)

        self.num_companies_selected_label = ttk.Label(self.master, text="0", font='SunValleyBodyStrongFont 18 bold')
        self.num_companies_selected_label.place(x=315,y=370,width=30,height=30)

        self.companies_label = ttk.Label(self.master, text="Companies Selected", font='SunValleyBodyStrongFont 15')
        self.companies_label.place(x=340,y=370,width=200, height=30)

        self.pyramid_title_string = tk.StringVar()
        self.pyramid_title_string.set(f"Copyright (2004) by Gentry Capital Corporation")

        #self.pyramid_title =  ttk.Entry(self.master, textvariable=self.pyramid_title_string)
        #self.pyramid_title.pack()

        self.customization_label = ttk.Label(self.master, text="Customization", font='SunValleyBodyStrongFont 12 bold')
        self.customization_label.place(x=250, y=20, width=150, height=20)

        self.options_label = ttk.Label(self.master, text="Options", font='SunValleyBodyStrongFont 12 bold')
        self.options_label.place(x=250, y=150, width=100, height=20)

        fname = open(path.join(user_data_dir(appname, appauthor), "default.txt")).read()
        if len(fname) >= 35:
            fname = f"...{fname[-30:]}"
        self.db_label = ttk.Label(self.master, text=f"Using file: {fname}")
        self.db_label.place(x=250,y=250, width=329, height=40)

        self.prepared_for_string = tk.StringVar()
        self.prepared_for_entry = ttk.Entry(self.master, textvariable=self.prepared_for_string)
        self.prepared_for_entry.place(x=380,y=50,width=195,height=30)

        self.prepared_for_label = ttk.Label(self.master, text="Prepared For")
        self.prepared_for_label.place(x=250,y=50,width=132,height=30)

        self.advisor_string = tk.StringVar()
        self.advisor_entry = ttk.Entry(self.master, textvariable=self.advisor_string)
        self.advisor_entry.place(x=380,y=90,width=195,height=30)
        self.advisor_label = ttk.Label(self.master, text="Advisor")
        self.advisor_label.place(x=250,y=90,width=132,height=30)

        self.generate_pyramid_var = tk.IntVar()
        self.generate_pyramid_var.set(1)
        self.generate_pyramid_checkbx = ttk.Checkbutton(self.master, text="Generate Pyramid Image", onvalue=1, offvalue=0, variable=self.generate_pyramid_var)
        self.generate_pyramid_checkbx.place(x=250,y=180, width=329, height=30)

        self.generate_word_doc = tk.IntVar()
        self.generate_word_doc.set(0)
        self.generate_word_doc_checkbx = ttk.Checkbutton(self.master, text="Generate Description Document", onvalue=1, offvalue=0, variable=self.generate_word_doc)
        self.generate_word_doc_checkbx.place(x=250,y=220, width=329, height=30)

        self.load_another_file_button = ttk.Button(self.master, text="Load Different Spreadsheet", command=self.load_another_file)
        self.load_another_file_button.place(x=250, y=290, width=325, height=30)

        self.create_pyramid_button = ttk.Button(self.master, text="Create", command=self.create_pyramid)
        #self.create_pyramid_button.pack()
        self.create_pyramid_button.place(x=250,y=450,width=150,height=30)

        self.reset_button = ttk.Button(self.master, text="Reset", command=self.reset_fields)
        self.reset_button.place(x=250+150+25, y=450, width=150, height=30)

    def load_another_file(self):
        print("Trying to load another file...")
        loaded_file = self.load_file()
        self.show_companies()

    def onselect(self, evt):
        w = evt.widget
        num_selected = len(w.curselection())
        self.num_companies_selected_label.config(text=f"{num_selected}")

    def reset_fields(self):
        self.generate_word_doc.set(0)
        self.generate_pyramid_var.set(1)
        self.prepared_for_string.set("")
        self.advisor_string.set("")
        self.company_listbox.listbox.selection_clear(0,'end')
        self.num_companies_selected_label.config(text="0")

    def create_pyramid(self):
        selected_companies_tmp = [self.company_listbox.listbox.get(index) for index in self.company_listbox.listbox.curselection()]
        print(f'Selected {len(selected_companies_tmp)} companies')
        selected_companies = [x.split(" - ")[0].strip() for x in selected_companies_tmp]
        print(f'Symbols: {len(selected_companies)}')
        selected_df = self.df[self.df["Symbol"].isin(selected_companies)].sort_values(by="Weighting", ascending=False)
        print(f'DF: {len(selected_df)}')
        #pyramid_file_path = asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx *.xls")])

        company_blurbs = {
        }

        if self.generate_word_doc.get() == 1:
            word_file_path = asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])

            if not word_file_path:
                return

            document = Document("assets/doc_template.docx")

            advisor = self.prepared_for_string.get()
            today = datetime.today()

            to_replace = document.paragraphs[1].text
            to_replace = to_replace.replace("{{ADVISOR}}", advisor)
            to_replace = to_replace.replace("{{DATE}}" , today.strftime("%B %m, %Y"))

            document.paragraphs[1].text = to_replace
            document.paragraphs[1].style.font.name = "Times New Roman"

            word_company_scores = {}

            for company, ticker, dividend, currency, desc, weighting in zip(selected_df["Company Name"], selected_df["Symbol"], selected_df["Dividend?"], selected_df["Currency"], selected_df["Blurb"], selected_df["Weighting"]):
                extra_char = ""
                if dividend.strip() == "Y":
                    extra_char = "***"
                if currency == "CAD":
                    canadian_companies.append(f'{company}{extra_char} ({ticker})')

                word_company_scores[f'{company}{extra_char} ({ticker})'] = weighting
                company_blurbs[f'{company}{extra_char} ({ticker})'] = desc

                # new_paragraph = document.add_paragraph()
                # temp_run = new_paragraph.add_run(f'{company}{extra_char} ({ticker})')
                # temp_run.bold = True
                # new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # document.paragraphs.pop()
                # desc_paragraph = document.add_paragraph(f"\n{desc}\n")
                # desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # desc_paragraph.bold = False

            word_company_scores = {k: v for k, v in sorted(word_company_scores.items(), key=lambda item: item[1])}
            word_pyramid = {}
            for company, score in word_company_scores.items():
                if score not in word_pyramid:
                    word_pyramid[score] = []
                word_pyramid[score].append(company)

            sorted_keys = sorted(word_pyramid.keys())

            companies = []


            for key in sorted_keys:
                if word_pyramid[key] == []:
                    continue
                if len(word_pyramid[key])<=15:
                    companies.append(word_pyramid[key])
                else:
                    temp_list = []
                    for x in range(len(word_pyramid[key])//15):
                        temp_list.append(word_pyramid[key][x*15:15])

                    temp_list.append(word_pyramid[key][::-1][:len(word_pyramid[key])%15])
                    for row in temp_list[::-1]:
                        if row != []:
                            companies.append(row)

            everything_list = []
            for row in companies:
                if row == []:
                    continue
                for company in row:
                    everything_list.append(company)

            companies = everything_list

            companies = pyramid_list(companies)

            everything_list = []

            for row in companies:
                for company in row:
                    new_paragraph = document.add_paragraph()
                    temp_run = new_paragraph.add_run(company)
                    temp_run.bold = True
                    temp_run.font.name = "Times New Roman"
                    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    document.paragraphs.pop()
                    desc_paragraph = document.add_paragraph(f"\n{company_blurbs[company]}\n\n\n")
                    desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    desc_paragraph.bold = False
                    desc_paragraph.style.font.name = "Times New Roman"

            document.save(word_file_path)

        if self.generate_pyramid_var.get() == 0:
            None
        else: 
            image_file_path = asksaveasfilename(defaultextension=".png", filetypes=[("Image files","*.png")])

            company_scores = {}

            # For each 'Company Name', company_scores[company] = 'Weighting' for that company
            for company, ticker, weighting, dividend, currency in zip(selected_df["Company Name"], selected_df["Symbol"], selected_df["Weighting"], selected_df["Dividend?"], selected_df["Currency"]):
                extra_char = ""
                if dividend.strip() == "Y":
                    extra_char = "*"
                if currency == "CAD":
                    canadian_companies.append(f'{company} ({ticker}){extra_char}')
                company_scores[f'{company} ({ticker}){extra_char}'] = weighting

            from PIL import Image, ImageDraw, ImageFont
            import math
            # Sort the dictionary by score in ascending order
            company_scores = {k: v for k, v in sorted(company_scores.items(), key=lambda item: item[1])}

            # Group the companies by their scores
            pyramid = {i: [] for i in range(1, 7)}
            for company, score in company_scores.items():
                pyramid[score].append(company)

            # Initialize some parameters
            img_width = 3300
            pyramid_width = 3000
            img_height = 1550 #2550
            overall_height = 2550
            block_color = (0, 150, 255)  # Blue color
            dividend_color = block_color #(1,50,32)  # Green color
            font_color = (255, 255, 255)  # White color
            normal_block = (0, 150, 255) # Dark Blue  
            padding = 15  # Padding around blocks
            radius = 10 
            logo_padding = 100 # top padding for logo
            pyramid_padding_bottom = 350 # bottom padding for pyramid
            max_companies_in_row = 10

            config = read_config(path.join(user_data_dir(appname, appauthor), "config.json"))
            print(config)
            img_width = config['img_width']
            pyramid_width = config['pyramid_width']
            img_height = config['img_height']
            overall_height = config['overall_height']
            block_color = tuple(config['block_color'])
            dividend_color = tuple(config['dividend_color'])
            font_color = tuple(config['font_color'])
            normal_block = tuple(config['normal_block'])
            padding = config['padding']
            radius = config['radius']
            logo_padding = config['logo_padding']
            pyramid_padding_bottom = config['pyramid_padding_bottom']
            max_companies_in_row = config['max_companies_in_row']

            base_path = "assets"
            primary_font = path.join(base_path, "baskerville.ttf")
            secondary_font = path.join(base_path, "gill_sans_bold.ttf")

            # Group the companies by their scores
            pyramid = {}
            for company, score in company_scores.items():
                if score not in pyramid:
                    pyramid[score] = []
                pyramid[score].append(company)

            # Sort the pyramid keys in ascending order to have a pyramid shape
            sorted_keys = sorted(pyramid.keys())

            companies = []


            for key in sorted_keys:
                if pyramid[key] == []:
                    continue
                if len(pyramid[key])<=max_companies_in_row:
                    companies.append(pyramid[key])
                else:
                    temp_list = []
                    for x in range(len(pyramid[key])//max_companies_in_row):
                        temp_list.append(pyramid[key][x*max_companies_in_row:max_companies_in_row])

                    temp_list.append(pyramid[key][::-1][:len(pyramid[key])%max_companies_in_row])
                    for row in temp_list[::-1]:
                        if row != []:
                            companies.append(row)

            everything_list = []
            for row in companies:
                if row == []:
                    continue
                for company in row:
                    everything_list.append(company)

            companies = everything_list
            print(f"Still have {len(companies)}")

            companies = pyramid_list(companies)

            for company_row in companies:
                print(company_row)

            # Calculate the maximum number of companies in a group (this will be the width of your pyramid)
            max_companies = max(len(row) for row in companies) #max(len(v) for v in pyramid.values())

            # Calculate the total number of groups (the height of your pyramid)
            total_groups = len(companies) #len(pyramid)

            # Calculate the size of each block based on the width and height of the image and the number of blocks
            block_size = min((pyramid_width - padding) // max_companies - padding, (img_height - padding) // total_groups - padding)

            # Calculate the size of each block based on the width and height of the image and the number of blocks
            block_width = (pyramid_width - padding) // max_companies - padding
            block_height = (img_height - padding) // total_groups - padding

            # Calculate the total width and height of the blocks (including padding)
            total_width = max_companies * (block_width + padding)
            total_height = total_groups * (block_height + padding)

            # Calculate the starting position for the first block
            start_x = (pyramid_width - total_width) // 2
            start_y = (img_height - total_height) // 2


            # Create an image big enough to hold the pyramid
            img = Image.new('RGB', (pyramid_width, img_height), "white")
            d = ImageDraw.Draw(img)

            def wrap_text(text, max_length):
                words = text.split()
                lines = []
                current_line = []

                for word in words:
                    if len(' '.join(current_line + [word])) <= max_length:
                        current_line.append(word)
                    else:
                        lines.append(' '.join(current_line))
                        current_line = [word]
                lines.append(' '.join(current_line))
                
                return '\n'.join(lines)

            min_font_size = 100_000_000

            for i, row in enumerate(companies):
                for j, company in enumerate(row):
                    font_size = min(block_width // (len(company) // 2 + 1), block_height // 2)
                    if font_size < min_font_size:
                        min_font_size = font_size

            num_dividends = 0

            # Loop over each level of the pyramid
            for i, row in enumerate(companies):
                for j, company in enumerate(row):
                    # Calculate the position of the block
                    x = start_x + j * (block_width + padding) + (max_companies - len(row)) * (block_width + padding) // 2
                    y = start_y + i * (block_size + padding)

                    # Calculate the color of the block
                    if company[-1] == "*":
                        block_color = dividend_color
                        num_dividends += 1
                    else:
                        block_color = normal_block

                    # Draw the block
                    d.rounded_rectangle([x, y, x + block_width, y + block_height], fill=block_color, radius=radius)

                    # Adjust font size based on the length of the company name and block size
                    font_size = min_font_size + 15
                    fnt = ImageFont.truetype(secondary_font, font_size)

                    # Implement word wrap for the company name
                    wrapped_company = wrap_text(company, (block_width // font_size)*1.5)

                    # Draw the company name
                    bbox = d.textbbox((x, y), wrapped_company, font=fnt)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]
                    text_x = x + (block_width - text_width) // 2
                    text_y = y + (block_height - text_height) // 2
                    if "Constellation" in company and len(row) == 5:
                        print("Manually shifting up")
                        text_y -= 25
                    d.text((text_x, text_y), wrapped_company, font=fnt, fill=font_color, align='center', spacing=10)


            print(image_file_path)

            # Save the image
            if not image_file_path:
                return

            image_logo = Image.open("./assets/GentryCapitalRGB.jpg")
            _, l_height = image_logo.size

            new_img = Image.new("RGB", (img_width, l_height), "white")
            left = (new_img.width - image_logo.width) // 2
            top = (new_img.height - image_logo.height) // 2

            new_img.paste(image_logo, (left, top))

            final_img = Image.new("RGB", (img_width, overall_height), "white")
            final_img.paste(new_img, (0, logo_padding))

            text1 = self.pyramid_title_string.get()
            text2 = f"({num_dividends} Dividend payors - all identified by asterisk)"
            text3 = "FOR INTERNAL USE ONLY"
            text4 = datetime.today().strftime("%B %d, %Y")
            #text5 = f"{len([self.company_listbox.listbox.get(index) for index in self.company_listbox.listbox.curselection()])}"
            text5 = f"Number of positions: {len([self.company_listbox.listbox.get(index) for index in self.company_listbox.listbox.curselection()])}"

            if self.prepared_for_string.get() != "":
                print(f"Prepared for string present - {self.prepared_for_string.get()}")
                if text5 != "":
                    if type(text5) == list:
                        text5.append(f"Prepared for {self.prepared_for_string.get()}")
                    else:
                        temp_prepared_for_text = f"Prepared for {self.prepared_for_string.get()}"
                        if len(temp_prepared_for_text) > 30:
                            text5 = [text5, "Prepared for: "]
                            for line in textwrap.wrap(self.prepared_for_string.get(), 30, break_long_words=False):
                                text5.append(line)
                        else:
                            text5 = [text5, f"Prepared for {self.prepared_for_string.get()}"]
            if self.advisor_string.get() != "":
                advisor = self.advisor_string.get()
                if text5 != "":
                    if type(text5) == list:
                        text5.append(f"By {advisor}")
                    else:
                        tmp_var = [text5, f"By {advisor}"]
                        text5 = tmp_var
                else:
                    text5 = f"By {advisor}"

            print(text5)
            draw = ImageDraw.Draw(final_img)
            font_size = 60

            font1 = ImageFont.truetype(primary_font, font_size)
            text_width1 = draw.textlength(text1, font=font1)
            while text_width1 > final_img.width:
                font_size -= 1
                font1 = ImageFont.truetype(primary_font, font_size)
                text_width1 = draw.textlength(text1, font=font1)

            font2_size = font_size - 10
            font2 = ImageFont.truetype(primary_font, font2_size)
            text_width2 = draw.textlength(text2, font=font2)
            while text_width2 > final_img.width:
                font_size -= 1
                font2 = ImageFont.truetype(primary_font, font2_size)
                text_width2 = draw.textlength(text2, font=font2)

            new_font_size = 30
            font3 = ImageFont.truetype(secondary_font, new_font_size)
            _,t,_,b = draw.textbbox((100,100), text3, font=font3)
            while (b-t) < (0.4*pyramid_padding_bottom):
                new_font_size += 1
                font3 = ImageFont.truetype(secondary_font, new_font_size)
                _,t,_,b = draw.textbbox((100,100), text3, font=font3)
            text_width3 = draw.textlength(text3, font=font3)
            while text_width3 > (final_img.width*0.8):
                new_font_size -= 1
                font3 = ImageFont.truetype(secondary_font, new_font_size)
                text_width3 = draw.textlength(text3, font=font3)
            _,t,_,b = draw.textbbox((100,100),text3,font=font3)

            font4_size = font2_size - 5
            font4 = ImageFont.truetype(primary_font, font4_size)
            text_width4 = draw.textlength(text4, font=font4)


            start_x1 = (final_img.width - text_width1) // 2
            start_y1 = image_logo.height + logo_padding + 50

            start_x2 = (final_img.width - text_width2) // 2
            start_y2 = image_logo.height + logo_padding + 125

            start_x3 = (final_img.width - text_width3) // 2
            start_y3 = (final_img.height - pyramid_padding_bottom + ((pyramid_padding_bottom)-(b-t))//2)

            start_x4 = (final_img.width - text_width4) - 150
            start_y4 = logo_padding + 30 

            start_x5 = 150
            start_y5 = start_y4
            
            draw.text((start_x1, start_y1), text1, font=font1, fill="black")
            draw.text((start_x2, start_y2), text2, font=font2, fill="black")
            draw.text((start_x3, start_y3), text3, font=font3, fill=(158,161,162))
            draw.text((start_x4, start_y4), text4, font=font4, fill="black")
            if type(text5) == str:
                draw.text((start_x5, start_y5), text5, font=font4, fill="black")
            else:
                for idx, text2write in enumerate(text5):
                    draw.text((start_x5, start_y5 + (idx*69)), text2write, font=font4, fill="black")

            final_img.paste(img, ((img_width-pyramid_width)//2, l_height + (overall_height - l_height - img_height) - pyramid_padding_bottom) )

            final_img.save(image_file_path)
            #final_img.show()
            #img.show()

            if platform.system() == 'Darwin':       # macOS
                subprocess.call(('open', image_file_path))
            elif platform.system() == 'Windows':    # Windows
                os.startfile(image_file_path)
            else:                                   # linux variants
                subprocess.call(('xdg-open', image_file_path))


if __name__ == "__main__":
    root = tk.Tk()
    wd = os.getcwd()
    try:
        os.chdir(sys._MEIPASS)
    except AttributeError:
        os.chdir(wd)
    print(os.listdir())
    app = CompanySelector(root)
    sv_ttk.set_theme("dark")
    if darkdetect.isLight():
        sv_ttk.set_theme("light")

    root.mainloop()
